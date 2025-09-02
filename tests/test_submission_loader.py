import os
from typing import Any

import pytest
from pytest_mock import MockerFixture

from gradebotguru.submission_loader import (
    load_code,
    load_docx,
    load_markdown,
    load_pdf,
    load_submission,
    load_text,
)


def test_load_text(tmp_path: Any) -> None:
    """
    Test loading a text file.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of the text file is correctly read.
    """
    text_file = tmp_path / "test.txt"
    text_file.write_text("This is a test text file.")
    assert load_text(str(text_file)) == "This is a test text file."


def test_load_markdown(tmp_path: Any) -> None:
    """
    Test loading a Markdown file.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of the Markdown file is correctly read.
    """
    md_file = tmp_path / "test.md"
    md_file.write_text("# Test\nThis is a test markdown file.")
    assert load_markdown(str(md_file)) == "# Test\nThis is a test markdown file."


def test_load_pdf(tmp_path: Any) -> None:
    """
    Test loading a PDF file.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of the PDF file is correctly read.
    """
    from pypdf import PdfWriter

    pdf_file = tmp_path / "test.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(pdf_file, "wb") as f:
        writer.write(f)
    assert load_pdf(str(pdf_file)) == ""


def test_load_docx(tmp_path: Any) -> None:
    """
    Test loading a DOCX file.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of the DOCX file is correctly read.
    """
    import docx

    docx_file = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(docx_file)
    assert load_docx(str(docx_file)) == "This is a test DOCX file."


def test_load_code(tmp_path: Any) -> None:
    """
    Test loading a code file.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of the code file is correctly read.
    """
    code_file = tmp_path / "test.py"
    code_file.write_text("print('Hello, world!')")
    assert load_code(str(code_file)) == "print('Hello, world!')"


def test_load_submission(tmp_path: Any) -> None:
    """
    Test loading different types of submission files.

    Parameters:
    - tmp_path (Any): A temporary directory provided by pytest.

    Asserts that the content of different types of files is correctly read using load_submission.
    """
    text_file = tmp_path / "test.txt"
    text_file.write_text("This is a test text file.")
    assert load_submission(str(text_file)) == "This is a test text file."

    md_file = tmp_path / "test.md"
    md_file.write_text("# Test\nThis is a test markdown file.")
    assert load_submission(str(md_file)) == "# Test\nThis is a test markdown file."

    from pypdf import PdfWriter

    pdf_file = tmp_path / "test.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(pdf_file, "wb") as f:
        writer.write(f)
    assert load_submission(str(pdf_file)) == ""

    import docx

    docx_file = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(docx_file)
    assert load_submission(str(docx_file)) == "This is a test DOCX file."

    code_file = tmp_path / "test.py"
    code_file.write_text("print('Hello, world!')")
    assert load_submission(str(code_file)) == "print('Hello, world!')"


def test_load_submissions(mocker: MockerFixture) -> None:
    """
    Test loading and parsing submissions from a directory.

    Mocks the presence of files in a directory and checks if `load_submissions`
    correctly reads and parses the files.

    Parameters:
    - mocker (MockerFixture): A pytest fixture for mocking.
    """
    mock_files: dict[str, str] = {
        "submission1.txt": "This is the content of submission 1.",
        "submission2.txt": "This is the content of submission 2.",
    }

    def mock_listdir(directory: str) -> list[str]:
        return list(mock_files.keys())

    def mock_isfile(path: str) -> bool:
        return path.split(os.sep)[-1] in mock_files

    def mock_open_func(filepath: str, mode: str = "r", encoding: str | None = None) -> Any:
        filename = filepath.split(os.sep)[-1]
        if filename in mock_files:
            return mocker.mock_open(read_data=mock_files[filename]).return_value
        raise FileNotFoundError

    mocker.patch("os.listdir", mock_listdir)
    mocker.patch("os.path.isfile", mock_isfile)
    mocker.patch("builtins.open", mock_open_func)

    submissions = {
        file: load_submission(f"path/to/submissions/{file}")
        for file in mock_listdir("path/to/submissions")
    }
    assert submissions == mock_files


if __name__ == "__main__":
    pytest.main()
