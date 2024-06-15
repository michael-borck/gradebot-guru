import os
from typing import Any, Dict

def load_text(file_path: str) -> str:
    """
    Load a text file and return its content as a string.

    Parameters:
    - file_path (str): The path to the text file.

    Returns:
    - str: The content of the text file.

    Examples:
    >>> from tempfile import NamedTemporaryFile
    >>> with NamedTemporaryFile('w', delete=False) as temp_file:
    ...     temp_file.write("This is a test text file.")
    >>> load_text(temp_file.name)
    'This is a test text file.'
    >>> os.remove(temp_file.name)
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def load_markdown(file_path: str) -> str:
    """
    Load a Markdown file and return its content as a string.

    Parameters:
    - file_path (str): The path to the Markdown file.

    Returns:
    - str: The content of the Markdown file.

    Examples:
    >>> from tempfile import NamedTemporaryFile
    >>> with NamedTemporaryFile('w', delete=False) as temp_file:
    ...     temp_file.write("# Test\\nThis is a test markdown file.")
    >>> load_markdown(temp_file.name)
    '# Test\\nThis is a test markdown file.'
    >>> os.remove(temp_file.name)
    """
    return load_text(file_path)

def load_pdf(file_path: str) -> str:
    """
    Load a PDF file and return its text content as a string.

    Parameters:
    - file_path (str): The path to the PDF file.

    Returns:
    - str: The text content of the PDF file.

    Requires:
    - pypdf package.

    Examples:
    >>> from pypdf import PdfWriter
    >>> from tempfile import NamedTemporaryFile
    >>> pdf_file = NamedTemporaryFile('wb', delete=False)
    >>> writer = PdfWriter()
    >>> writer.add_blank_page(width=72, height=72)
    >>> writer.write(pdf_file.name)
    >>> pdf_file.close()
    >>> load_pdf(pdf_file.name)
    ''
    >>> os.remove(pdf_file.name)
    """
    try:
        from pypdf import PdfReader
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            return ' '.join(page.extract_text() for page in reader.pages)
    except ImportError:
        raise ImportError("pypdf is required to load PDF files")

def load_docx(file_path: str) -> str:
    """
    Load a DOCX file and return its text content as a string.

    Parameters:
    - file_path (str): The path to the DOCX file.

    Returns:
    - str: The text content of the DOCX file.

    Requires:
    - python-docx package.

    Examples:
    >>> import docx
    >>> from tempfile import NamedTemporaryFile
    >>> docx_file = NamedTemporaryFile(delete=False)
    >>> doc = docx.Document()
    >>> _ = doc.add_paragraph("This is a test DOCX file.")
    >>> doc.save(docx_file.name)
    >>> load_docx(docx_file.name)
    'This is a test DOCX file.'
    >>> os.remove(docx_file.name)
    """
    try:
        import docx
        doc = docx.Document(file_path)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)
    except ImportError:
        raise ImportError("python-docx is required to load DOCX files")

def load_code(file_path: str) -> str:
    """
    Load a code file and return its content as a string.

    Parameters:
    - file_path (str): The path to the code file.

    Returns:
    - str: The content of the code file.

    Examples:
    >>> from tempfile import NamedTemporaryFile
    >>> with NamedTemporaryFile('w', delete=False) as temp_file:
    ...     temp_file.write("print('Hello, world!')")
    >>> load_code(temp_file.name)
    "print('Hello, world!')"
    >>> os.remove(temp_file.name)
    """
    return load_text(file_path)

def load_submission(file_path: str) -> str:
    """
    Load a submission file based on its extension and return its content as a string.

    Parameters:
    - file_path (str): The path to the submission file.

    Returns:
    - str: The content of the submission file.

    Raises:
    - ValueError: If the file format is not supported.

    Examples:
    >>> from tempfile import NamedTemporaryFile
    >>> with NamedTemporaryFile('w', delete=False, suffix='.txt') as temp_file:
    ...     temp_file.write("This is a test text file.")
    >>> load_submission(temp_file.name)
    'This is a test text file.'
    >>> os.remove(temp_file.name)

    >>> with NamedTemporaryFile('w', delete=False, suffix='.md') as temp_file:
    ...     temp_file.write("# Test\\nThis is a test markdown file.")
    >>> load_submission(temp_file.name)
    '# Test\\nThis is a test markdown file.'
    >>> os.remove(temp_file.name)

    >>> with NamedTemporaryFile('w', delete=False, suffix='.py') as temp_file:
    ...     temp_file.write("print('Hello, world!')")
    >>> load_submission(temp_file.name)
    "print('Hello, world!')"
    >>> os.remove(temp_file.name)

    >>> with NamedTemporaryFile('w', delete=False, suffix='.unsupported') as temp_file:
    ...     temp_file.write("Unsupported content.")
    >>> try:
    ...     load_submission(temp_file.name)
    ... except ValueError as e:
    ...     str(e)
    'Unsupported file format: .unsupported'
    >>> os.remove(temp_file.name)
    """
    _, file_extension = os.path.splitext(file_path)
    loaders = {
        '.txt': load_text,
        '.md': load_markdown,
        '.pdf': load_pdf,
        '.docx': load_docx,
        '.py': load_code,
        '.js': load_code,
        '.css': load_code,
        '.html': load_code,
    }
    loader = loaders.get(file_extension.lower())
    if loader:
        return loader(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
